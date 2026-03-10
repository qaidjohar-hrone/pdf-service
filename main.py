import io
import logging
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict, Optional

import docx
import jinja2
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from weasyprint import CSS, HTML

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Document Template Service",
    description="API for generating PDFs from Word/HTML templates with variable substitution",
    version="1.0.0",
)


# Request models
class DocumentRequest(BaseModel):
    template_name: str = Field(
        ..., description="Name of the template file (without extension)"
    )
    variables: Dict[str, Any] = Field(
        ..., description="Variables to substitute in the template"
    )

    class Config:
        json_schema_extra = {
            "example": {
                "template_name": "invoice_template",
                "variables": {
                    "customer_name": "John Doe",
                    "invoice_number": "INV-2024-001",
                    "amount": 1500.00,
                },
            }
        }


class HTMLRequest(BaseModel):
    template_name: str = Field(
        ..., description="Name of the HTML template file (without extension)"
    )
    variables: Dict[str, Any] = Field(
        ..., description="Variables to substitute in the template"
    )

    class Config:
        json_schema_extra = {
            "example": {
                "template_name": "report_template",
                "variables": {
                    "title": "Monthly Report",
                    "date": "2024-01-15",
                    "content": "Report details here",
                },
            }
        }


# Configuration
TEMPLATES_DIR = Path("templates")
WORD_TEMPLATES_DIR = TEMPLATES_DIR / "word"
HTML_TEMPLATES_DIR = TEMPLATES_DIR / "html"
OUTPUT_DIR = Path("output")

# Create directories if they don't exist
for directory in [WORD_TEMPLATES_DIR, HTML_TEMPLATES_DIR, OUTPUT_DIR]:
    directory.mkdir(parents=True, exist_ok=True)

# Initialize Jinja2 for HTML templates
jinja_env = jinja2.Environment(
    loader=jinja2.FileSystemLoader(str(HTML_TEMPLATES_DIR)),
    autoescape=jinja2.select_autoescape(["html", "xml"]),
)


def replace_variables_in_word(doc_path: Path, variables: Dict[str, Any]) -> Path:
    """
    Replace variables in Word document using {{variable_name}} syntax.
    Supports both paragraph and table cell content.
    """
    try:
        doc = docx.Document(doc_path)

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in variables.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    # Handle runs to preserve formatting
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in variables.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                for run in paragraph.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(
                                            placeholder, str(value)
                                        )

        # Save modified document
        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_docx.name)
        return Path(temp_docx.name)

    except Exception as e:
        logger.error(f"Error replacing variables in Word document: {e}")
        raise HTTPException(
            status_code=500, detail=f"Failed to process Word document: {str(e)}"
        )


def convert_word_to_pdf_libreoffice(docx_path: Path) -> bytes:
    """
    Convert Word document to PDF using LibreOffice (works on Linux/Mac/Windows).
    This is the recommended method for cross-platform compatibility.
    """
    try:
        # Create temporary directory for output
        temp_dir = tempfile.mkdtemp()

        # Run LibreOffice conversion
        # --headless: run without GUI
        # --convert-to pdf: output format
        # --outdir: output directory
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                temp_dir,
                str(docx_path),
            ],
            capture_output=True,
            text=True,
            timeout=30,
        )

        if result.returncode != 0:
            logger.error(f"LibreOffice conversion failed: {result.stderr}")
            raise Exception(f"LibreOffice conversion failed: {result.stderr}")

        # Find the generated PDF
        pdf_path = Path(temp_dir) / f"{docx_path.stem}.pdf"

        if not pdf_path.exists():
            raise Exception("PDF file was not created")

        # Read PDF content
        with open(pdf_path, "rb") as f:
            pdf_content = f.read()

        # Cleanup
        import shutil

        shutil.rmtree(temp_dir)

        return pdf_content

    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=500, detail="PDF conversion timed out")
    except FileNotFoundError:
        raise HTTPException(
            status_code=500,
            detail="LibreOffice not found. Please install: sudo apt-get install libreoffice",
        )
    except Exception as e:
        logger.error(f"Error converting Word to PDF: {e}")
        raise HTTPException(
            status_code=500, detail=f"Failed to convert to PDF: {str(e)}"
        )


def render_html_to_pdf(template_name: str, variables: Dict[str, Any]) -> bytes:
    """Render HTML template with variables and convert to PDF using WeasyPrint."""
    try:
        # Load and render template
        template = jinja_env.get_template(f"{template_name}.html")
        html_content = template.render(**variables)

        # Convert HTML to PDF using WeasyPrint
        # Fixed: Use string parameter correctly
        pdf_bytes = HTML(string=html_content).write_pdf()

        return pdf_bytes

    except jinja2.TemplateNotFound:
        raise HTTPException(
            status_code=404, detail=f"Template '{template_name}.html' not found"
        )
    except Exception as e:
        logger.error(f"Error rendering HTML to PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")


@app.post("/api/v1/generate-pdf/word", response_class=StreamingResponse)
async def generate_pdf_from_word(request: DocumentRequest):
    """
    Generate PDF from Word document template.

    - Reads Word document from templates/word/{template_name}.docx
    - Replaces variables in format {{variable_name}}
    - Converts to PDF using LibreOffice and returns as download

    Note: Requires LibreOffice installed (sudo apt-get install libreoffice)
    """
    template_path = WORD_TEMPLATES_DIR / f"{request.template_name}.docx"

    if not template_path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"Template '{request.template_name}.docx' not found in templates/word/",
        )

    try:
        # Process document
        processed_docx = replace_variables_in_word(template_path, request.variables)

        # Convert to PDF using LibreOffice
        pdf_content = convert_word_to_pdf_libreoffice(processed_docx)

        # store in output directory (optional)
        output_pdf_path = OUTPUT_DIR / f"{request.template_name}.pdf"
        with open(output_pdf_path, "wb") as f:
            f.write(pdf_content)

        # Cleanup temporary docx
        os.unlink(processed_docx)

        # Return PDF as streaming response
        return StreamingResponse(
            io.BytesIO(pdf_content),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={request.template_name}.pdf"
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in generate_pdf_from_word: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/v1/generate-pdf/html", response_class=StreamingResponse)
async def generate_pdf_from_html(request: HTMLRequest):
    """
    Generate PDF from HTML template.

    - Reads HTML template from templates/html/{template_name}.html
    - Renders with Jinja2 using provided variables
    - Converts to PDF using WeasyPrint and returns as download
    """
    try:
        # Render and convert
        pdf_content = render_html_to_pdf(request.template_name, request.variables)

        # store in output directory (optional)
        output_pdf_path = OUTPUT_DIR / f"{request.template_name}.pdf"
        with open(output_pdf_path, "wb") as f:
            f.write(pdf_content)

        # Return PDF as streaming response
        return StreamingResponse(
            io.BytesIO(pdf_content),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={request.template_name}.pdf"
            },
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in generate_pdf_from_html: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/v1/templates/word")
async def list_word_templates():
    """List all available Word templates."""
    templates = [f.stem for f in WORD_TEMPLATES_DIR.glob("*.docx")]
    return {"templates": templates}


@app.get("/api/v1/templates/html")
async def list_html_templates():
    """List all available HTML templates."""
    templates = [f.stem for f in HTML_TEMPLATES_DIR.glob("*.html")]
    return {"templates": templates}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
